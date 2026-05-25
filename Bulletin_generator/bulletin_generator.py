from docx import Document
import pandas as pd
from docx.shared import Pt          # Pt - форматування тексту
from docx.oxml.ns import qn         # Форматування таблиці
from docx.oxml import OxmlElement   # Форматування таблиці


# pd.set_option('display.max_columns', None)  
# pd.set_option('display.max_rows', None)
# pd.set_option('display.width', None)

date = '4 травня'
week = '18'

docx_template_path = "bulletin_template.docx"
main_db_path = "../main_database/flu.xlsx"
pcr_info_path = "../../../Disk D/Мои документы/Grip/Rezultat/Зразки на грип та ГРВІ 2025 – 2026.xlsx"


df_pcr_list = pd.read_excel(pcr_info_path, header=None)
df_main = pd.read_excel(main_db_path)
df_flu = df_main[df_main['line'].isin(['B', 'H1N1', 'H3N2'])]


total_PCR_quantity = len(df_pcr_list)
total_sequenced = len(df_main)
flu_sequenced = len(df_flu)


def get_flu_zrazok_or_zrazkiv(total_sequenced):
    flu_zrazok_or_zrazkiv_dict = {
        1: 'зразок',
        2: 'зразки',
        3: 'зразки',
        4: 'зразки'
    }

    if 11 <= total_sequenced % 100 <= 14:
        return'зразків'
    else:
        return flu_zrazok_or_zrazkiv_dict.get(total_PCR_quantity % 10, 'зразків')

PCR_total_zrazok_or_zrazkiv = get_flu_zrazok_or_zrazkiv(total_PCR_quantity)
total_zrazok_or_zrazkiv = get_flu_zrazok_or_zrazkiv(total_sequenced)
flu_zrazok_or_zrazkiv = get_flu_zrazok_or_zrazkiv(flu_sequenced)


value_count_flu_series = df_flu['region'].value_counts()

kyiv_flu_quantity = value_count_flu_series["Київський МЦКПХ"]
value_count_flu_series = value_count_flu_series.drop("Київський МЦКПХ")



region_rename = {
    'Вінницький ОЦКПХ': 'Вінницької',
    'Волинський ОЦКПХ': 'Волинської',
    'Дніпропетровський ОЦКПХ': 'Дніпропетровської',
    'Житомирський ОЦКПХ': 'Житомирської',
    'Закарпатський ОЦКПХ': 'Закарпатської',
    'Запорізький ОЦКПХ': 'Запорізької',
    'Ів.-Франківський ОЦКПХ': 'Івано-Франківської',
    'Київський ОЦКПХ': 'Київської',
    'Кіровоградський ОЦКПХ': 'Кіровоградської',
    'Львівський ОЦКПХ': 'Львівської',
    'Миколаївський ОЦКПХ': 'Миколаївської',
    'Одеський ОЦКПХ': 'Одеської',
    'Полтавський ОЦКПХ': 'Полтавської',
    'Рівненський ОЦКПХ': 'Рівненської',
    'Сумський ОЦКПХ': 'Сумської',
    'Тернопільський ОЦКПХ': 'Тернопільської',
    'Харківський ОЦКПХ': 'Харківської',
    'Херсонський ОЦКПХ': 'Херсонської',
    'Хмельницький ОЦКПХ': 'Хмельницької',
    'Черкаський ОЦКПХ': 'Черкаської',
    'Чернівецький ОЦКПХ': 'Чернівецької',
    'Чернігівський ОЦКПХ': 'Чернігівської'
}

value_count_flu_series.index = value_count_flu_series.index.map(region_rename)

regions_paragraph = ""
for oblast, value in zip(value_count_flu_series.index, value_count_flu_series):
    regions_paragraph += f'{oblast} – {value}, ' 
    
# Remove ', ' from the last item
regions_paragraph = regions_paragraph[:-2]


# flu_table = df_flu[['reference', 'line', 'clade', 'subclade']].value_counts()
flu_table = df_flu[['line', 'reference', 'clade', 'subclade']].value_counts()
flu_table_df = flu_table.reset_index()

flu_table_df.columns = ['Субтип/лінія', 'Референс', 'Клада', 'Субклада', 'Кількість']

flu_table_df = flu_table_df.sort_values(
    by=['Субтип/лінія', 'Кількість'], 
    ascending=[False, False]
)




def insert_table_after_placeholder(doc, placeholder, df):

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            
            paragraph.text = "Таблиця 1. Генетична характеристика вірусів грипу"
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            
            # Створюємо таблицю: к-ть рядків (дані + заголовок) і к-ть колонок
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Normal Table'
            table.autofit = False

            # Чорна сітка
            tblBorders = OxmlElement('w:tblBorders')
            borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
            for border_name in borders:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')   # Тип лінії - суцільна
                border.set(qn('w:space'), '0')      # Прибирає відступ міжд словом і лінією
                tblBorders.append(border)
            table._tbl.tblPr.append(tblBorders)


            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            table._tbl.tblPr.append(tblLayout)

            # Фіксуємо ширину макета
            column_widths = [Pt(90), Pt(172), Pt(119), Pt(71), Pt(72)]

            # 1. Задаємо ширину для колонок
            for j, width in enumerate(column_widths):
                table.columns[j].width = width

            # 2. Жорстко фіксуємо ЗАГАЛЬНУ ширину таблиці
            total_width = sum(w.twips for w in column_widths)
            tblW = table._tbl.tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblW.set(qn('w:type'), 'dxa')
                tblW.set(qn('w:w'), str(int(total_width)))

            
            # Зменшення відступу таблиці зліва
            tblInd = OxmlElement('w:tblInd')
            indent_value = Pt(-30)
            tblInd.set(qn('w:w'), str(int(indent_value.twips))) 
            tblInd.set(qn('w:type'), 'dxa')
            table._tbl.tblPr.append(tblInd)


            # Заповнюємо заголовки
            for j, column_name in enumerate(df.columns):
                cell = table.cell(0, j)
                cell.text = column_name
                cell.width = column_widths[j]
            

            # Заповнюємо дані
            for i, row in enumerate(df.values):
                for j, value in enumerate(row):
                    cell = table.cell(i + 1, j)
                    cell.text = str(value)       
                    cell.width = column_widths[j]
            

            # Форматуємо шрифт у всій таблиці під Times New Roman 14
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)
            return


doc = Document(docx_template_path)

# Функція для заміни тексту у всіх елементах документа
def replace_text_in_document(doc, search_text, replace_text):
    # Збираємо всі параграфи (з тексту та таблиць)
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for p in paragraphs:
        if search_text in p.text:
            # Зберігаємо оновлений текст
            new_text = p.text.replace(search_text, replace_text)
            
            # Очищуємо всі існуючі фрагменти (runs)
            # Це важливо, щоб Word не залишав "сміття" від старих стилів
            for run in p.runs:
                run.text = ""
            
            # Створюємо один новий run з усім текстом
            new_run = p.add_run(new_text)
            
            # Застосовуємо Times New Roman 14
            new_run.font.name = 'Times New Roman'
            new_run.font.size = Pt(14)
            
            # Спеціальний фікс для кирилиці (щоб не злітало в Calibri)
            r = new_run._element.get_or_add_rPr().get_or_add_rFonts()
            r.set(qn('w:ascii'), 'Times New Roman')
            r.set(qn('w:hAnsi'), 'Times New Roman')
            r.set(qn('w:cs'), 'Times New Roman')
            r.set(qn('w:eastAsia'), 'Times New Roman')





    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             if search_text in cell.text:
    #                 cell.text = cell.text.replace(search_text, replace_text)

replacements = {
    "{day_month}": date,
    "{total_PCR_quantity}": str(total_PCR_quantity),
    "{PCR_total_zrazok_or_zrazkiv}": PCR_total_zrazok_or_zrazkiv,
    "{total_sequenced}": str(total_sequenced),
    "{total_zrazok_or_zrazkiv}": total_zrazok_or_zrazkiv,
    "{flu_sequenced}": str(flu_sequenced),
    "{flu_zrazok_or_zrazkiv}": flu_zrazok_or_zrazkiv,
    "{kyiv_flu_quantity}": str(kyiv_flu_quantity),
    "{regions_paragraph}": regions_paragraph
}


for key, value in replacements.items():
    replace_text_in_document(doc, key, value)

insert_table_after_placeholder(doc, "{flu_table}", flu_table_df)

doc.save(f"Бюлетень {week} тиждень.docx")

